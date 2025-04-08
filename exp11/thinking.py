from docx import Document
from docx.shared import Pt

# 创建一个新的Word文档
document = Document()

# 添加标题
document.add_heading('会议通知', 0)

# 添加表格相关信息
table_info = [
    ["会议名称：", ""],
    ["会议时间：", ""],
    ["会议地点：", ""],
    ["会议主持：", ""],
    ["出席人员：", ""]
]

# 先添加信息部分的表格
table1 = document.add_table(rows=len(table_info), cols=2)
for i, (label, value) in enumerate(table_info):
    cell1 = table1.cell(i, 0)
    cell2 = table1.cell(i, 1)
    cell1.text = label
    cell2.text = value

# 定义会议议程表格的表头
agenda_headers = ["内容", "汇报人", "专题列席人", "汇报时间"]
# 添加会议议程表格
table2 = document.add_table(rows=8, cols=len(agenda_headers))
# 设置表头
for i, header in enumerate(agenda_headers):
    cell = table2.cell(0, i)
    cell.text = header
    # 可以设置表头字体等格式（可选）
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].runs[0].font.bold = True

# 添加备注信息
document.add_paragraph("注：会议重要，请提前五分钟到达会场。会议期间，请将手机置于震动档。")

# 添加联系人信息（这里假设联系人信息为示例值，实际可替换）
document.add_paragraph("联系人：张三")

# 保存文档
document.save('meeting_notice.docx')
